from flask import Flask, request, jsonify
from docx import Document

app = Flask(__name__)

@app.route('/api/employees')
def getAllEmployees():
    return {'employees': [{'id':1, 'name':'Balram'},{'id':2, 'name':'Tom'}]}

@app.route('/api/document', methods=['POST'])
def update_document():
    requestData = request.get_json()

    sourcePath = requestData['sourcePath']
    searchTextPara = requestData['searchStringPara']
    replaceTextPara = requestData['replacementStringPara']
    searchTextTable = requestData['searchStringTable']
    replaceTextTable = requestData['replacementStringTable']
    targetElement = requestData['targetElement']
    targetPath = requestData['targetPath']
    performSave = False

    document = Document(sourcePath)
    elementIndex = targetElement.index('para') if 'para' in targetElement else -1

    # if 'para' in targetElement:
    if elementIndex > -1:
        for index, para in enumerate(document.paragraphs):
            if searchTextPara in para.text:
                para.runs[0].text = para.runs[0].text.replace(searchTextPara, replaceTextPara)
                performSave = True

    elementIndex = targetElement.index('table') if 'table' in targetElement else -1
    # if 'table' in targetElement:
    if elementIndex > -1:
        for index,table in enumerate(document.tables):
            for colIndex, column in enumerate(table.columns):
                for cellIndex, cell in enumerate(column.cells):
                    if (cellIndex > 0):
                        if searchTextTable in cell.text:
                            cell.text = cell.text.replace(searchTextTable, replaceTextTable)
                            performSave = True

    if performSave:
        document.save(targetPath)
        return jsonify('saved successfully at ' + targetPath)


if __name__ == '__main__':
     app.run(port=5000)
